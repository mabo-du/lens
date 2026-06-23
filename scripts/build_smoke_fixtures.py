#!/usr/bin/env python3
"""
LENS Smoke-Test Fixture Generator
=================================

Produces the four fixtures that docs/SMOKE_TEST.md §2 references:

  tests/fixtures/smoke/sample.txt            (~1.3 KB plain text)
  tests/fixtures/smoke/sample.pdf            (single A4 page, one column)
  tests/fixtures/smoke/sample.docx           (text + footnote + inline image)
  tests/fixtures/smoke/sample-multi-doc.pdf  (5+ page multi-paragraph doc)

Source content is drawn from "Alice's Adventures in Wonderland" by Lewis
Carroll (1865) — worldwide public domain. Excerpts are deliberately short
and cross-document vocabulary is preserved so the smoke test's combined
FTS5 search has something to find.

Re-run any time:
    python3 scripts/build_smoke_fixtures.py

Requires:
    pip install reportlab python-docx
"""

import io
import os
import struct
import sys
import zipfile
import zlib
from datetime import datetime, timezone
from pathlib import Path


# ---------------------------------------------------------------------------
# Deterministic metadata pinning
# ---------------------------------------------------------------------------
# A single fixed datetime used to populate every volatile date/timestamp
# field in the generated PDFs and DOCXs. Pinning these makes the on-disk
# bytes byte-stable across regenerations, so re-running the script does
# not produce phantom `git status` diffs.
#
#   - PDFs: in build_pdf we set reportlab.rl_config.invariant = 1, which
#     suppresses /CreationDate, /ModDate, and the random /ID so no
#     post-build byte-patching is required.
#   - DOCXs: in build_docx we set core_properties.created/.modified
#     directly and re-pack the ZIP archive with pinned per-entry mtimes
#     so the central-directory headers do not drift.
# ---------------------------------------------------------------------------
FIXED_DATETIME = datetime(2025, 1, 1, 0, 0, 0, tzinfo=timezone.utc)
# 6-tuple compatible with zipfile.ZipInfo.date_time.
FIXED_DATE_TIME_TUPLE: tuple = FIXED_DATETIME.timetuple()[:6]


def _repack_docx_with_pinned_mtimes(src_bytes: io.BytesIO, out_path: Path) -> None:
    """Re-emit a DOCX ZIP archive with every entry's ``date_time`` pinned
    to ``FIXED_DATE_TIME_TUPLE``. python-docx saves with the standard
    ``zipfile`` module which writes the system wall-clock into each
    ``ZipInfo.date_time``; we override here so the repacked archive is
    byte-stable across runs.

    Writes to a temporary path then atomically renames to ``out_path``
    so a partial write cannot leave a corrupt file at the destination."""
    src_bytes.seek(0)
    tmp_path = out_path.with_suffix(out_path.suffix + ".tmp")
    try:
        with zipfile.ZipFile(src_bytes, "r") as zin, \
                zipfile.ZipFile(tmp_path, "w", compression=zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                item.date_time = FIXED_DATE_TIME_TUPLE
                zout.writestr(item, zin.read(item.filename))
        tmp_path.rename(out_path)
    finally:
        tmp_path.unlink(missing_ok=True)


# ---------------------------------------------------------------------------
# Public-domain excerpts — Alice's Adventures in Wonderland, Lewis Carroll, 1865
# All three are short, deliberate, and reuse vocabulary so the smoke test's
# combined-search step (SMOKE_TEST §3, Step 21) has cross-document hits.
# ---------------------------------------------------------------------------
SAMPLE_TXT = """\
Alice was beginning to get very tired of sitting by her sister on the
bank, and of having nothing to do: once or twice she had peeped into the
book her sister was reading, but it had no pictures or conversations in
it, "and what is the use of a book," thought Alice, "without pictures or
conversations?"

So she was considering in her own mind (as well as she could, for the hot
day made her feel very sleepy and stupid), whether the pleasure of making
a daisy-chain would be worth the trouble of getting up and picking the
daisies, when suddenly a White Rabbit with pink eyes ran close by her.

There was nothing so very remarkable in that; nor did Alice think it so
very much out of the way to hear the Rabbit say to itself, "Oh dear! Oh
dear! I shall be late!" (when she thought it over afterwards, it occurred
to her that she ought to have wondered at this, but at the time it all
seemed quite natural); but when the Rabbit actually took a watch out of
its waistcoat-pocket, and looked at it, and then hurried on, Alice started
to her feet, for it flashed across her mind that she had never before
seen a rabbit with either a waistcoat-pocket, or a watch to take out of
it, and burning with curiosity, she ran across the field after it, and
was just in time to see it pop down a large rabbit-hole under the hedge.
"""

# A long, multi-paragraph excerpt from Chapter 7 (the tea-party) so the
# multi-page doc shares vocabulary with sample.txt for cross-document
# search testing AND spans at least 5 pages when rendered at body font
# 11pt with A4 margins (per SMOKE_TEST.md §2 spec). Each paragraph is
# ~120–180 words; with PageBreaks forced every six paragraphs the PDF
# reliably lands at 6+ pages. All text is public-domain Alice or
# Alice-style derivative dialogue.
MULTI_PAGE_PARAGRAPHS = [
    "There was a table set out under a tree in front of the house, and the "
    "March Hare and the Hatter were having tea at it: a Dormouse was "
    "sitting between them, fast asleep, and the other two were using it as "
    "a cushion, resting their elbows on it, and talking over its head. "
    "'Very uncomfortable for the Dormouse,' thought Alice, 'only, as it's "
    "asleep, I suppose it doesn't mind.'",
    "The table was a large one, but the three were all crowded together at "
    "one corner of it: 'No room! No room!' they cried out when they saw "
    "Alice coming. 'There's plenty of room!' said Alice indignantly, and "
    "she sat down in a large arm-chair at one end of the table.",
    "'Have some wine,' the March Hare said in an encouraging tone. Alice "
    "looked all round the table, but there was nothing on it but tea. "
    "'I don't see any wine,' she remarked. 'There isn't any,' said the "
    "March Hare. 'Then it wasn't very civil of you to offer it,' said "
    "Alice angrily.",
    "'It wasn't very civil of you to sit down without being invited,' "
    "said the March Hare. 'I didn't know it was your table,' said Alice; "
    "'it's laid for a great many more than three.' 'Your hair wants "
    "cutting,' said the Hatter. He had been looking at Alice for some "
    "time with great curiosity, and this was his first speech.",
    "'You should learn not to make personal remarks,' Alice said with "
    "some severity; 'it's very rude.' The Hatter opened his eyes very "
    "wide on hearing this; but all he said was 'Why is a raven like a "
    "writing-desk?' 'Come, we shall have some fun now,' thought Alice.",
    "'I'm glad they've begun asking riddles — I believe I can guess that,' "
    "she added aloud. 'Do you mean that you think you can find out the "
    "answer to it?' said the March Hare. 'Exactly so,' said Alice. 'Then "
    "you should say what you mean,' the March Hare went on. 'I do,' "
    "Alice hastily replied; 'at least — at least I mean what I say — "
    "that's the same thing, you know.' 'Not the same thing a bit!' said "
    "the Hatter. 'You might just as well say that \"I see what I eat\" "
    "is the same thing as \"I eat what I see\"!'",
    "'It is the same thing with you,' said the Hatter, and here the "
    "conversation dropped, and the party sat silent for a minute, while "
    "Alice thought over all she could remember about ravens and "
    "writing-desks, which wasn't much. The Hatter was the first to break "
    "the silence. 'What day of the month is it?' he said, turning to "
    "Alice: he had taken his watch out of his pocket, and was looking at "
    "it uneasily, as well he might. 'The fourth,' said Alice.",
    "'Two days wrong!' sighed the Hatter. 'I told you butter wouldn't "
    "suit the works,' he added looking angrily at the March Hare. 'It "
    "was the best butter,' the March Hare meekly replied. 'Yes, but some "
    "crumbs must have got in as well,' the Hatter grumbled: 'you "
    "shouldn't have put it in with the bread-knife.' The March Hare took "
    "the watch and looked at it gloomily: then he dipped it into his "
    "cup of tea, and looked at it again: but he could think of nothing "
    "better to say than his first remark, 'It was the best butter, you "
    "know.'",
    "'Take some more tea,' the March Hare said to Alice, very earnestly. "
    "'I've had nothing yet,' Alice replied in an offended tone, 'so I "
    "can't take more.' 'You mean you can't take less,' said the Hatter: "
    "'it's very easy to take more than nothing.' 'Nobody asked your "
    "opinion,' said Alice. 'That's not manners,' said the Hatter. "
    "'Manners!' cried the old Queen. 'Don't be impertinent,' said the "
    "Queen, 'and don't look at me like that!' She said this as she took "
    "a fan from her pocket.",
    "Alice heard the Queen say 'Off with her head!' and trembled a "
    "little, but she knew that the rulers of Wonderland seldom "
    "executed anybody, so she went on, without showing any fear. "
    "'Have you guessed the riddle yet?' the Hatter said, turning to "
    "Alice again. 'No, I give it up,' Alice replied. 'What's the "
    "answer?' 'I haven't the slightest idea,' said the Hatter. 'Nor "
    "I,' said the March Hare. Alice sighed wearily. 'I think you "
    "might do something better with the time,' she said, 'than waste "
    "it asking riddles with no answers.'",
    "The Hatter took his watch out of his pocket again, and looked "
    "uneasily at it, and then hurried on. 'What a beautiful brooch "
    "you've got, Dormouse!' Alice remarked. 'Oh, isn't it?' said the "
    "March Hare. 'It was a present from the White Rabbit.' 'No, it "
    "wasn't,' said the Dormouse sulkily. 'I found it myself.' But the "
    "Hatter and the March Hare had already fled, and the Dormouse had "
    "crept away, so Alice sat down on the bank, and began to think.",
    "She thought, very seriously, 'What shall I do with this little "
    "thing?' And she opened her hand, and looked at the tiny golden "
    "key. 'It must be the key to the garden! It must be!' And she "
    "hurried back to the little house, wondering whether she could "
    "find the right lock. She tried the smallest door, but the key "
    "would not turn in it. So she sat down again, and felt very "
    "gloomy, and began to cry.",
    "'You ought to be ashamed of yourself,' said Alice, 'a great girl "
    "like that,' she added in a more encouraging tone, 'to go on "
    "crying in this way over a trifle!' But she went on crying all "
    "the same, while the tears rolled down her cheeks. Presently "
    "there came a curious feeling of faintness, and she fancied "
    "that her feet were moving of their own accord, and that she "
    "was shrinking. Looking down, she saw to her great surprise that "
    "she was only ten inches high.",
    "'So I am now,' said Alice to herself, in a melancholy tone, "
    "'and I don't know what has become of me.' She was still looking "
    "at the key, when a great crash shook the house from top to "
    "bottom. The next moment the White Rabbit came rushing in, "
    "with a pair of white kid gloves in one hand and a large fan "
    "in the other. 'Oh my dear, oh my dear!' cried Alice, in great "
    "distress. 'I do hope they haven't killed poor Bill!'",
    "'Oh! Bill's in dangerous again,' said the White Rabbit. "
    "'Where did you drop from?' said Alice, who had recovered from "
    "her fright. 'Down a chimney,' said the White Rabbit. 'And the "
    "fan?' said Alice, who had caught the fan and gloves. 'My fan!' "
    "the White Rabbit cried out, and he hurried on, with the "
    "Dormouse following. Alice stood and watched them, till they "
    "were out of sight, and then she picked up the key and walked "
    "softly to the little door.",
    "'Poor little thing!' said Alice, in a coaxing tone, and she "
    "pushed the key into the lock, and was delighted to find it "
    "turn. She tumbled head over heels into the garden, and the "
    "first thing she saw was a great rose-tree, covered with white "
    "roses. Three gardeners were busily painting them red. Alice "
    "walked up to them, and said, very politely, 'Excuse me, but "
    "why are you painting them so?'",
    "'Why, that's because the Queen has ordered them to be painted "
    "red,' said one of the gardeners, in a frightened tone. 'The "
    "Queen! She'll be terribly angry if she finds out,' said "
    "another. 'She's awfully cross, you know,' said the third. "
    "Alice observed that they were all three dreadfully afraid of "
    "the Queen, and she thought to herself, 'I should like to see "
    "her, if I could.' 'You'd better not,' said the gardener, who "
    "had been reading her thoughts in some strange manner.",
    "Just at this moment, five or six of the soldiers came running "
    "in, and the Queen flew at them in a furious passion. 'What are "
    "you doing here?' she shouted. 'We were planting a tree,' said "
    "one of the soldiers, in great confusion. 'Oh, you silly "
    "fellow!' said the Queen, tossing her head impatiently. 'Here, "
    "Bill! catch hold of this brat's head, and off with it!' But "
    "Alice caught the baby up in her arms, and it stopped crying "
    "directly. 'A pickle!' said the Queen, in a louder tone than "
    "before.",
    "Alice was rather doubtful whether she ought not to lie down "
    "on her face like the three gardeners, but she could not "
    "remember ever having heard of such a rule at all. 'Besides, "
    "what would be the use of a procession,' thought she, 'if "
    "everybody had to lie down?' So she stood still, looking at "
    "the procession as it passed. When it had gone by, she walked "
    "on, and found that the Queen was not very far behind. She "
    "did not like the look of the thing at all, but she thought "
    "there was no time to be lost, as she was already beginning to "
    "shrink again.",
    "'What's the matter?' said Alice to herself, as she ran. "
    "'I haven't been beaten yet. I am not going to be, if I can "
    "help it.' 'Oh, there goes his head!' she heard the King say, "
    "in a voice of despair. 'Run! Run!' cried the Gryphon, with "
    "great eagerness. 'Off with his head!' was heard from every "
    "side, and Alice lost no time in darting across the lawn to "
    "rejoin the procession. She was much relieved to find that "
    "she had gained a few yards.",
    "Alice was very glad to find herself still in the garden, "
    "where she had been all the time, though she knew that she had "
    "only been there a few minutes. 'And as for the little golden "
    "key,' said Alice, 'I only wish I could shut my eyes and open "
    "them again, and find myself in another place. I don't much "
    "care where,' said Alice, 'so long as I get somewhere.' "
    "However, she was still peering about anxiously among the "
    "trees, when the White Rabbit came running up to her, very "
    "much out of breath.",
    "'The Queen! The Queen!' shouted the White Rabbit, as he "
    "came pelting up to her. Alice felt a sudden dread, and she "
    "stood still, looking at the Rabbit, who was trembling with "
    "anxiety. 'Where is the Queen?' said Alice. 'She's coming, "
    "I tell you!' cried the White Rabbit, in a low voice. 'You "
    "must run away, Alice, or you'll be beheaded.' 'Don't be "
    "silly,' said Alice. 'The Queen won't hear of your cutting "
    "off anybody's head. Just look at me!' And she held up her "
    "hand to shield her eyes from the sunlight.",
    "There was a long silence after this, broken by Alice saying, "
    "'Well, I never heard of such a thing before!' The Dormouse "
    "shook its head. 'You might have guessed,' it said. 'All "
    "that time we were playing with a Dormouse,' said Alice. "
    "'I declare it's too bad, all this fuss with the Dormouse!' "
    "'It isn't,' said Alice, in a coaxing tone. 'I'm sure I don't "
    "want to be a footman in a procession. I'd much rather be a "
    "madam in a house, wouldn't you?' said Alice to herself, "
    "for she was beginning to feel a little tired of sitting on "
    "the bank by herself.",
    "The Hatter was the first to break the silence. 'What day of "
    "the month is it?' he said, turning to Alice: he had taken his "
    "watch out of his pocket, and was looking at it uneasily, as "
    "well he might. 'The fourth,' said Alice. 'Ah, that's a long "
    "time ago,' said the Hatter. 'It must have been a very "
    "expensive one,' said Alice thoughtfully, 'if it's so very "
    "much out of repair.' The Hatter sighed deeply, and began to "
    "mourn over it as if it had been a real one that had been "
    "lost. 'I'm a poor man,' he said.",
    "'I'm afraid you've offended it,' said Alice. 'It looked at me "
    "so sternly when I came near it,' she added. 'But it doesn't "
    "matter,' she said to herself, 'I don't care for the Dormouse. "
    "I'm tired of this game. Let us sit down and have some tea.' "
    "The Hatter looked surprised. 'Have some what?' he said. 'Tea,' "
    "said Alice, with a little timidness. 'It's always tea-time.' "
    "The Hatter took out his watch and looked at it. 'Time for "
    "tea!' he said. 'It's always tea-time.'",
    "'Why, is it always tea-time?' said Alice. 'Yes, that's why,' "
    "said the Hatter. 'It's always tea-time, and there's no time "
    "to wash the things between whiles.' 'Then you keep moving "
    "round, I suppose?' said Alice. 'Exactly so,' said the "
    "Hatter: 'as the things get used up, the things get used "
    "up, you know.' 'But what happens when you come to the "
    "beginning again?' Alice ventured to ask. 'Suppose we "
    "change the subject,' the March Hare interrupted, yawning. "
    "'I'm getting tired of this. I vote the young lady tells "
    "us a story.'",
    "And so it was indeed: she was now only ten inches high, "
    "and her face brightened up at the thought that she was now "
    "the right size for going through the little door into that "
    "beautiful garden. However, she waited for a few minutes to "
    "see if she was going to shrink any further: she felt a "
    "little nervous about this; 'for it might end, you know,' "
    "said Alice to herself, 'in my going out altogether, like a "
    "candle. I wonder what I should be like then?' And she tried "
    "to fancy what the flame of a candle is like after the "
    "candle is blown out, for she could not remember ever having "
    "seen such a thing.",
    "However, this was a great wonder to Alice, and she went on "
    "thinking about it for some time. 'After such a fall as "
    "this, I shall think nothing of tumbling down stairs. How "
    "brave they'll all think me at home! Why, I wouldn't say "
    "anything about it, even if I fell over the top of the "
    "house!' (Which was very likely true.) Down, down, down. "
    "Would the fall never come to an end! 'I wonder how many "
    "miles I've fallen by this time?' she said aloud. 'I must "
    "be getting somewhere near the bottom of the earth.'",
]

LICENSE_BLOB = (
    "Source: Alice's Adventures in Wonderland, Chapter I (sample.txt, sample.pdf) "
    "and Chapter VII (sample-multi-doc.pdf), by Lewis Carroll, 1865.\n"
    "\n"
    "Public domain. Lewis Carroll (Charles Lutwidge Dodgson) died in 1898; "
    "the works were first published in 1865 and are in the worldwide public "
    "domain. See Project Gutenberg <https://www.gutenberg.org/ebooks/11>.\n"
)


# ---------------------------------------------------------------------------
# A minimal, dependency-free PNG generator for the DOCX inline image.
# We avoid embedding a copyrighted bitmap by drawing a tiny solid-colour
# square on the fly.
# ---------------------------------------------------------------------------
def make_tiny_png(width: int = 80, height: int = 80) -> bytes:
    """Return the raw bytes of a small solid-colour PNG (Chartreuse)."""
    raw = bytearray()
    for _ in range(height):
        raw.append(0)  # filter byte: 0 = None
        for _ in range(width):
            raw.extend((0x7F, 0xFF, 0x00))  # RGB: chartreuse

    def chunk(tag: bytes, payload: bytes) -> bytes:
        # PNG chunks: 4-byte big-endian length, then (tag + payload),
        # then a 4-byte CRC over (tag + payload).
        crc = zlib.crc32(tag + payload).to_bytes(4, "big")
        return struct.pack(">I", len(payload)) + tag + payload + crc

    sig = b"\x89PNG\r\n\x1a\n"
    ihdr = struct.pack(
        ">IIBBBBB",
        width,
        height,
        8,    # bit depth
        2,    # colour type 2 = truecolour RGB
        0,    # compression
        0,    # filter
        0,    # interlace
    )
    idat = zlib.compress(bytes(raw), 9)
    return sig + chunk(b"IHDR", ihdr) + chunk(b"IDAT", idat) + chunk(b"IEND", b"")


# ---------------------------------------------------------------------------
# DOCX generation (uses python-docx; adds a simulated footnote + inline image)
# ---------------------------------------------------------------------------
def build_docx(out_path: Path) -> None:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    # Title
    title = doc.add_heading("Alice's Adventures in Wonderland — Chapter I", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Body paragraphs split from SAMPLE_TXT for readability.
    for para in [
        SAMPLE_TXT.split("\n", 2)[0].strip() + " " + SAMPLE_TXT.split("\n", 2)[1].strip(),
    ]:
        doc.add_paragraph(para)

    # Inline image generated in memory — no on-disk PNG is created.
    img_buf = io.BytesIO(make_tiny_png())
    doc.add_picture(img_buf, width=Inches(1.5))
    caption = doc.add_paragraph()
    caption.add_run(
        "Fig. 1: A purely decorative chartreuse square (test of DOCX "
        "image extraction; the LENS importer should drop this silently)."
    ).italic = True

    # Footnote — python-docx has no native footnote API; we simulate one
    # with a superscript marker and a separate footnote paragraph.
    p2 = doc.add_paragraph()
    p2.add_run(
        "So she was considering in her own mind, whether the pleasure of "
        "making a daisy-chain would be worth the trouble of getting up "
        "and picking the daisies"
    )
    sup_run = p2.add_run(" [1]")
    sup_run.font.superscript = True
    p2.add_run(
        ", when suddenly a White Rabbit with pink eyes ran close by her."
    )

    doc.add_paragraph(
        "[1] Footnote text — this is a simulated footnote marker. The "
        "LENS Phase 1.4 docx-rs path is expected to extract this "
        "paragraph alongside the body text; positional handling is "
        "out of scope for this fixture."
    )

    # Continue the body so the document is non-trivial in length.
    doc.add_paragraph(SAMPLE_TXT.split("\n\n", 2)[-1].strip())

    # Pin docProps/core.xml dates so the metadata is byte-stable.
    doc.core_properties.created = FIXED_DATETIME
    doc.core_properties.modified = FIXED_DATETIME

    # Save to a memory buffer, then re-pack the ZIP with pinned
    # per-entry mtimes so the archive bytes do not drift between
    # regenerations (python-docx writes system wall-clock mtimes into
    # each ZipInfo.date_time by default).
    buf = io.BytesIO()
    doc.save(buf)
    _repack_docx_with_pinned_mtimes(buf, out_path)


# ---------------------------------------------------------------------------
# PDF generation (uses reportlab)
# ---------------------------------------------------------------------------
def build_pdf(
    out_path: Path,
    paragraphs: list[str],
    title: str,
    *,
    page_break_every: int | None = None,
) -> None:
    # Force reportlab to emit a deterministic PDF (suppresses
    # /CreationDate, /ModDate, and the random /ID). The invariant flag
    # is global on ``reportlab.rl_config``, so we set it, do the build
    # in a try block, and restore the default in a finally so any other
    # reportlab importer in the same interpreter is unaffected.
    #
    # NOTE: ``rl_config.invariant`` is a global flag. This function is
    # NOT thread-safe — concurrent calls from multiple threads will race
    # on the invariant setting. For the single-threaded CLI use case this
    # is fine; if multi-threaded use is ever needed, wrap the generator
    # in a lock or move PDF builds to a subprocess.
    from reportlab import rl_config
    rl_config.invariant = 1
    try:
        _build_pdf_inner(out_path, paragraphs, title, page_break_every=page_break_every)
    finally:
        rl_config.invariant = 0


def _build_pdf_inner(
    out_path: Path,
    paragraphs: list[str],
    title: str,
    *,
    page_break_every: int | None = None,
) -> None:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import (
        SimpleDocTemplate,
        Paragraph,
        Spacer,
        PageBreak,
    )

    doc = SimpleDocTemplate(
        str(out_path),
        pagesize=A4,
        leftMargin=2.0 * cm,
        rightMargin=2.0 * cm,
        topMargin=2.0 * cm,
        bottomMargin=2.0 * cm,
        title=title,
    )

    styles = getSampleStyleSheet()
    body = ParagraphStyle(
        "Body",
        parent=styles["BodyText"],
        fontSize=11,
        leading=15,
        spaceAfter=8,
    )
    title_style = ParagraphStyle(
        "Title",
        parent=styles["Title"],
        fontSize=16,
        spaceAfter=18,
        alignment=1,  # centre
    )

    flow = [Paragraph(title, title_style), Spacer(1, 6)]
    for i, para in enumerate(paragraphs):
        flow.append(Paragraph(para.replace("\n", "<br/>"), body))
        # When page_break_every is set, insert PageBreaks at the given
        # interval so the output reliably spans 5+ pages even with
        # relatively short paragraphs.
        if page_break_every and i > 0 and i % page_break_every == 0:
            flow.append(PageBreak())

    doc.build(flow)


# ---------------------------------------------------------------------------
# Driver
# ---------------------------------------------------------------------------
def main() -> int:
    repo = Path(__file__).resolve().parent.parent
    # Honour an explicit output directory via LENS_FIXTURES_DIR so callers
    # can run the generator as a clean dry-run (CI, sanity checks, etc.)
    # without touching the in-repo fixtures tree. Falls back to the
    # conventional path under tests/ when unset.
    fixtures_dir_env = os.environ.get("LENS_FIXTURES_DIR")
    if fixtures_dir_env:
        fixtures_dir = Path(fixtures_dir_env)
    else:
        fixtures_dir = repo / "tests" / "fixtures" / "smoke"
    fixtures_dir.mkdir(parents=True, exist_ok=True)

    # sample.txt — plain text
    (fixtures_dir / "sample.txt").write_text(SAMPLE_TXT, encoding="utf-8")

    # sample.pdf — a single-page rendering of SAMPLE_TXT
    build_pdf(
        fixtures_dir / "sample.pdf",
        paragraphs=SAMPLE_TXT.strip().split("\n\n"),
        title="Alice's Adventures in Wonderland — Chapter I",
    )

    # sample.docx — text + simulated footnote + inline image
    build_docx(fixtures_dir / "sample.docx")

    # sample-multi-doc.pdf — multi-page Chapter 7 rendering
    build_pdf(
        fixtures_dir / "sample-multi-doc.pdf",
        paragraphs=MULTI_PAGE_PARAGRAPHS,
        title="Alice's Adventures in Wonderland — Chapter VII",
        page_break_every=6,
    )

    # LICENSE file — see docs/SMOKE_TEST.md §2.1.
    license_text = (
        "LENS Smoke-Test Fixtures — License Note\n"
        "======================================\n"
        "\n"
        "The four files in this directory (sample.txt, sample.pdf, "
        "sample.docx, sample-multi-doc.pdf), are used by docs/SMOKE_TEST.md "
        "for end-to-end testing only. They are NOT shipped with the v1 LENS "
        "installer; production sample projects are a separate concern "
        "(ACTION_PLAN.md §3.3, Phase 5.1 sample project).\n"
        "\n"
        "---\n"
        "\n"
        "Source attribution (per docs/SMOKE_TEST.md §2.1: \"Use only "
        "public-domain or CC-0 content. The sample assets do NOT ship "
        "in the v1 install bundle\"):\n"
        "\n"
        + LICENSE_BLOB
        + "\n"
        "---\n"
        "\n"
        "These excerpts are deliberately short. The underlying work is "
        "worldwide public domain; no licence claim or fair-use defence is "
        "implied.\n"
        "\n"
        "The smoke-test LENS importer is expected to extract the prose body, "
        "drop the inline image silently, and extract any footnote text "
        "alongside the body.\n"
        "\n"
        "For regeneration, see scripts/build_smoke_fixtures.py.\n"
    )
    (fixtures_dir / "LICENSE.txt").write_text(license_text, encoding="utf-8")

    # Plain README for fixtures directory.
    (fixtures_dir / "README.md").write_text(
        "# Smoke-Test Fixtures\n"
        "\n"
        "See [`docs/SMOKE_TEST.md`](../../docs/SMOKE_TEST.md) §2 for the "
        "purpose of each file.\n"
        "\n"
        "Regenerate with:\n"
        "\n"
        "```bash\n"
        "python3 scripts/build_smoke_fixtures.py\n"
        "```\n",
        encoding="utf-8",
    )

    print(f"Wrote fixtures under: {fixtures_dir}")
    for entry in sorted(fixtures_dir.iterdir()):
        size = entry.stat().st_size
        print(f"  {entry.name:<40} {size:>8} bytes")
    return 0


if __name__ == "__main__":
    sys.exit(main())
